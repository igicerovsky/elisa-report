""" Final markdown report assembly
"""
import re
import pandas as pd
from docx import Document
from docx.shared import Pt

from elisarep.typing import PathLikeOrNone
from .constants import CV_DIGITS
from .config import config as cfg
from .config import LIMITS_NAME, SOP_NAME, MHF_NAME


SHEET_FONT_SZ = Pt(8)


def plate_section_ex(df, plate):
    """ Plate section assembly to markdown
    """
    # df is formatted!
    md = f'### Plate {plate}\n\n'

    dg = CV_DIGITS
    md += df.to_markdown(floatfmt=f'#.{dg}f')
    md += '\n\n'
    md += '\* sample will be retested\n\n'

    return md


def assembly(reports, protocol, **kwargs):
    """ Whole report assembly to markdown
    """
    sop = cfg[SOP_NAME]
    md = f'# GT Analytics - Capsid {protocol}\n\n{sop}\n\n'

    md += '## Objective\n\n'
    md += 'Capsid concentration is determined for unknown samples.  \n\n'

    md += '## Method Status\n\n'
    mhf = cfg[MHF_NAME]
    md += f'For detailed method status see either {sop} and/or method history file {mhf}.  \n\n'

    md += '## Results - Current Reference\n\n'

    for r in reports:
        md += plate_section_ex(r['df'], r['plate'])

    md += '## Evaluation criteria\n\n'
    limits = cfg[LIMITS_NAME]
    md += (f'Validity of the assay: Intermediary control sample limits (3s) are '
           f'{limits[0]:.3e} - {limits[1]:.3e} cp/ml.  \n\n')
    if kwargs and kwargs['comment']:
        md += '## Comments\n\n'
        md += 'Add comment here...\n'

    return md


def check_bold(val: str):
    """ Check if value shall be bold
    """
    ret = None
    r = re.compile(r'(\*\*).*?(\*\*)')
    if r.match(val):
        ret = re.sub(r'[\*\*]', '', val)
    return ret


def check_retest(val: str):
    """ Check if value shall be retested
    """
    if ')*' in val:
        return True
    return None


def plate_sheet(df: pd.DataFrame, table):
    """ Generate sheet for given plate
    """
    # access first row's cells
    heading_row = table.rows[0].cells

    # add headings
    heading_row[0].text = 'Sample type'
    for i, h in enumerate(df.columns):
        heading_row[i + 1].text = h

    # index
    for i, r in enumerate(df.index):
        table.rows[i+1].cells[0].text = r

    retest = None
    for i in range(len(df.index)):
        for j in range(len(df.columns)):
            val = str(df.iat[i, j])
            cell = table.rows[i + 1].cells[j + 1]
            check_val = check_bold(val)
            if not retest:
                retest = check_retest(val)
            if check_val:
                val = check_val
            cell.text = val
            if check_val:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.bold = True

    # set table font to smaller size to fit into the table
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = SHEET_FONT_SZ

    return retest


def new_doc(reference_doc: PathLikeOrNone = None):
    """ Create new Word document
    """
    document = Document(reference_doc)
    for paragraph in document.paragraphs:
        if len(paragraph.text) == 0:
            p = paragraph._element
            p.getparent().remove(p)
            p._p = p._element = None

    return document


def assembly_word(reports: list, protocol: str, **kwargs):
    """ Word document final report assembly
    """
    reference_doc = None
    if kwargs and kwargs['reference_doc']:
        reference_doc = kwargs['reference_doc']
    document = new_doc(reference_doc)

    sop = cfg[SOP_NAME]
    document.add_heading(f'GT Analytics - Capsid {protocol}', 0)
    document.add_paragraph(f'{sop}')

    h = document.add_heading('Objective', level=1)
    h.style = document.styles['Heading 1']
    document.add_paragraph(
        'Capsid concentration is determined for unknown samples.')

    document.add_heading('Method Status', level=1)
    mhf = cfg[MHF_NAME]
    document.add_paragraph(
        f'For detailed method status see either {sop} and/or method history file {mhf}.')

    document.add_heading('Results - Current Reference', level=1)

    # add results for each plate
    for r in reports:
        document.add_heading(f'Plate {r["plate"]}', level=2)
        df = r['df']
        table = document.add_table(
            rows=len(df.index)+1, cols=len(df.columns)+1, style='Table Grid')
        retest = plate_sheet(df, table)
        if retest:
            p = document.add_paragraph('* sample will be retested')
            p.runs[0].font.size = SHEET_FONT_SZ
        if r != reports[-1]:
            document.add_page_break()

    document.add_heading('Evaluation criteria', level=1)
    limits = cfg[LIMITS_NAME]
    document.add_paragraph((f'Validity of the assay: '
                            f'Intermediary control sample limits (3s) are '
                            f'{limits[0]:.3e} - {limits[1]:.3e} cp/ml.')
                           )

    if kwargs and kwargs['docx_path']:
        document.save(kwargs['docx_path'])

    return document
